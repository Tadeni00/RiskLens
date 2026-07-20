"""
FraudTrap — Generate Comprehensive Study Guide (Word Document)
Produces a professional .docx file covering all platform components.
"""
from __future__ import annotations
import subprocess
import sys
from pathlib import Path


def ensure_python_docx():
    try:
        import docx
    except ImportError:
        print("[*] python-docx not found — installing via pip …")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx


def build_document():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn

    doc = Document()

    # ── Styles ─────────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    toc_style = doc.styles.add_style("TOCHeading", WD_STYLE_TYPE.PARAGRAPH)
    toc_style.font.name = "Calibri"
    toc_style.font.size = Pt(16)
    toc_style.font.bold = True
    toc_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    def add_bullet(text: str, level: int = 0):
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25 * level)
        return p

    def add_code_block(lines: list[str]):
        for line in lines:
            p = doc.add_paragraph(line)
            p.style = doc.styles["Normal"]
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def add_table(headers: list[str], rows: list[list[str]]):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = val
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)

    # ═══════════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("FraudTrap")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("Complete Study Guide")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x4A, 0x6F, 0x8C)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("Real-Time Fraud Detection for African Banks & Fintechs")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver_p.add_run("Version 2.0")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph("Table of Contents", style="TOCHeading")
    doc.add_paragraph()

    toc_entries = [
        ("Chapter 1", "Introduction"),
        ("Chapter 2", "Architecture Overview"),
        ("Chapter 3", "Data Model"),
        ("Chapter 4", "Scoring Pipeline"),
        ("Chapter 5", "Rules Engine"),
        ("Chapter 6", "Machine Learning Models"),
        ("Chapter 7", "Behavioral Intelligence"),
        ("Chapter 8", "Dashboard"),
        ("Chapter 9", "API Reference"),
        ("Chapter 10", "Deployment"),
        ("Chapter 11", "Operations"),
        ("Appendix A", "Configuration Reference"),
    ]
    for num, title in toc_entries:
        p = doc.add_paragraph()
        run = p.add_run(f"{num}:  {title}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 1 — INTRODUCTION
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Chapter 1: Introduction", level=1)

    doc.add_heading("1.1 What is FraudTrap?", level=2)
    doc.add_paragraph(
        "FraudTrap is a production-grade, real-time fraud detection platform designed "
        "for banks, fintechs, and financial institutions operating across Africa. "
        "It provides a scoring API that accepts transaction payloads, returns risk scores "
        "and decisions (APPROVE / REVIEW / BLOCK), stores recent decisions for dashboard "
        "consumption, accepts fraud labels from chargeback and manual review systems, "
        "and includes a Streamlit dashboard plus a live traffic simulator for continuous "
        "monitoring and development."
    )
    doc.add_paragraph(
        "The platform implements a three-phase machine learning lifecycle that evolves "
        "from unsupervised anomaly detection (no labels required) to fully supervised "
        "fraud classification as labeled data accumulates over time. This design is "
        "critical for African markets where labeled fraud data is initially scarce."
    )

    doc.add_heading("1.2 Target Users", level=2)
    add_bullet("Banks — Traditional banking institutions seeking real-time fraud prevention")
    add_bullet("Fintechs — Digital-first financial services providers (mobile money, neobanks)")
    add_bullet("Payment Processors — Card networks, payment gateways, and switch providers")
    add_bullet("Regulators — Compliance officers requiring audit trails and explainability")

    doc.add_heading("1.3 Key Differentiators", level=2)
    add_bullet("African Market Focus — Designed for African payment corridors, mobile money ecosystems, and regulatory environments (CBN, CBK, FSCA)")
    add_bullet("Multi-Tenant Architecture — Complete tenant isolation for banks and fintechs sharing the same platform instance")
    add_bullet("Three-Phase ML Lifecycle — Cold Start (unsupervised) → Semi-Supervised → Supervised, eliminating the cold-start problem")
    add_bullet("Behavioral Intelligence — Real-time customer, merchant, device, and beneficiary behavioral profiles with Redis hot storage")
    add_bullet("Explainability — SHAP-based explanations for supervised models, component breakdowns for cold-start, rule weight contributions")
    add_bullet("Sub-100ms SLA — Production scoring with P95 latency under 90ms")
    add_bullet("Champion-Challenger Architecture — Automated model evaluation and promotion with rollback support")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 2 — ARCHITECTURE OVERVIEW
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Chapter 2: Architecture Overview", level=1)

    doc.add_heading("2.1 System Components", level=2)
    add_table(
        ["Component", "Technology", "Role"],
        [
            ["Scoring API", "FastAPI (Python)", "Primary integration point; accepts transactions, returns risk scores"],
            ["Dashboard", "Streamlit", "Real-time monitoring, EDA, drift detection, compliance views"],
            ["Scoring Engine", "Python + NumPy", "Feature assembly, rule evaluation, model inference, policy floor"],
            ["ML Models", "PyTorch, scikit-learn, CatBoost, XGBoost, LightGBM", "Fraud detection across three lifecycle phases"],
            ["Event Backbone", "Apache Kafka", "Transactional event streaming for scores, labels, audit, drift alerts"],
            ["Online Feature Store", "Redis 7.2", "Velocity counters, behavioral profiles, recent score cache, blocklists"],
            ["Analytics Store", "ClickHouse 24.4", "Analytical rollups, drift metrics, historical scoring data"],
            ["Metadata Store", "PostgreSQL 16", "Model registry mirror, label persistence, training datasets"],
            ["Experiment Tracking", "MLflow", "Training runs, metrics, model versioning, artifact storage"],
        ],
    )

    doc.add_heading("2.2 Data Flow (Text Description)", level=2)
    doc.add_paragraph(
        "When a transaction arrives at the /v1/score endpoint, it flows through the following pipeline:"
    )
    add_code_block([
        "1. HTTP Request arrives at FastAPI /v1/score",
        "2. Pydantic validates the TransactionRequest payload",
        "3. Feature Assembly: Redis pipelines fetch velocity counters, behavioral profiles,",
        "   and entity history (~5-15ms)",
        "4. Rules Engine (Tier 1): Declarative rules evaluate blocklists, thresholds,",
        "   expressions, velocity, and geo rules (<1ms)",
        "5. If HARD BLOCK → return immediately with score=1.0, decision=BLOCK",
        "6. Otherwise → continue to ML scoring:",
        "   a. Behavioral Engine generates velocity, trust, similarity, novelty features",
        "   b. Model Orchestrator selects phase-appropriate model",
        "   c. Model inference (10-50ms)",
        "   d. Policy floor: max(model_score, heuristic_score)",
        "   e. Soft rule boost: if triggered rules, add risk_boost (capped at 0.30)",
        "7. Decision mapping: <0.40→APPROVE, 0.40-0.85→REVIEW, >=0.85→BLOCK",
        "8. Async: GNN scoring (non-blocking), SHAP explanation (if REVIEW band)",
        "9. Audit emit: Kafka + Redis recent_scores cache",
        "10. Return ScoringResponse to client",
    ])

    doc.add_heading("2.3 Deployment Architecture", level=2)
    doc.add_paragraph(
        "FraudTrap is deployed via Docker Compose with the following service topology:"
    )
    add_table(
        ["Service", "Port", "Health Check"],
        [
            ["API", "8000", "curl -f http://localhost:8000/health"],
            ["Dashboard", "8501", "Streamlit built-in health"],
            ["Redis", "6379", "redis-cli ping"],
            ["Kafka", "9092", "kafka-broker-api-versions"],
            ["Zookeeper", "2181", "nc -z localhost 2181"],
            ["ClickHouse", "9000/8123", "wget --spider http://localhost:8123/ping"],
            ["PostgreSQL", "5432", "pg_isready -U fraudtrap"],
            ["MLflow", "5000", "HTTP health check"],
        ],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 3 — DATA MODEL
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Chapter 3: Data Model", level=1)

    doc.add_heading("3.1 TransactionRequest Schema", level=2)
    doc.add_paragraph(
        "The inbound scoring payload from client banks and fintechs. "
        "Mandatory fields are the absolute minimum for real-time scoring. "
        "Optional fields improve accuracy — clients are encouraged to send all they have."
    )

    doc.add_heading("Mandatory Fields", level=3)
    add_table(
        ["Field", "Type", "Description"],
        [
            ["transaction_id", "str", "Unique transaction identifier (idempotency key); auto-generated UUID if omitted"],
            ["tenant_id", "str", "Client bank / fintech identifier (e.g., bank_ng_gtb)"],
            ["account_id", "str", "Tokenised account identifier"],
            ["amount", "float", "Transaction amount in local currency (must be > 0)"],
            ["currency", "str", "ISO 4217 currency code (3 chars, auto-uppercased)"],
            ["timestamp", "datetime", "Transaction initiation time (UTC)"],
            ["transaction_type", "str", "PAYMENT, TRANSFER, WITHDRAWAL, TOP_UP, REFUND, LOAN_DISBURSEMENT"],
            ["channel", "str", "WEB, MOBILE, API, POS, ATM, USSD"],
        ],
    )

    doc.add_paragraph()
    doc.add_heading("Optional Fields", level=3)
    add_table(
        ["Field", "Type", "Description"],
        [
            ["session_id", "str", "Tokenised session identifier"],
            ["merchant_id", "str", "Tokenised merchant identifier"],
            ["merchant_category_code", "str", "MCC code (e.g., 5411 for grocery)"],
            ["merchant_country", "str", "ISO 3166-1 alpha-2 country code"],
            ["counterparty_account_id", "str", "Tokenised destination account for transfers"],
            ["device_id", "str", "Tokenised device fingerprint"],
            ["device_type", "str", "MOBILE, DESKTOP, TABLET, POS_TERMINAL"],
            ["ip_address_hash", "str", "Hashed IP address (privacy-safe)"],
            ["user_agent_hash", "str", "Hashed user-agent string"],
            ["latitude", "float", "Geolocation latitude (-90 to 90)"],
            ["longitude", "float", "Geolocation longitude (-180 to 180)"],
            ["country_code", "str", "ISO 3166-1 alpha-2 country code"],
            ["typing_cadence_ms", "float", "Mean inter-keystroke interval in ms (from SDK)"],
            ["session_duration_seconds", "float", "Session duration in seconds"],
            ["field_visit_count", "int", "Number of form fields visited"],
            ["extra_fields", "dict", "Pass-through dict for client-specific fields"],
        ],
    )

    doc.add_heading("3.2 ScoringResponse Schema", level=2)
    add_table(
        ["Field", "Type", "Description"],
        [
            ["transaction_id", "str", "Echoed transaction identifier"],
            ["tenant_id", "str", "Tenant identifier"],
            ["risk_score", "float", "Fraud probability 0.0–1.0"],
            ["decision", "str", "APPROVE, REVIEW, or BLOCK"],
            ["model_phase", "str", "UNSUPERVISED, SEMI_SUPERVISED, SUPERVISED, or RULES"],
            ["model_version", "str", "Version of the model used for scoring"],
            ["latency_ms", "float", "Total scoring latency in milliseconds"],
            ["explanation", "Explanation?", "SHAP explanation (if available)"],
            ["triggered_rules", "list[str]", "IDs of rules that fired"],
            ["trace_id", "str", "Unique trace ID for audit log lookup"],
            ["scored_at", "datetime", "Timestamp when scored (UTC)"],
        ],
    )

    doc.add_heading("3.3 LabelPayload Schema", level=2)
    doc.add_paragraph(
        "Ground-truth labels arriving from chargeback systems or manual review. "
        "Labels are consumed by the Kafka label worker and used to train models."
    )
    add_table(
        ["Field", "Type", "Description"],
        [
            ["transaction_id", "str", "Transaction being labelled"],
            ["tenant_id", "str", "Tenant identifier"],
            ["label", "int", "1 = fraud, 0 = legitimate"],
            ["label_source", "str", "CHARGEBACK, MANUAL_REVIEW, or DISPUTE_RESOLVED"],
            ["chargeback_reason_code", "str?", "Visa/MC reason code (used to filter non-fraud chargebacks)"],
            ["labelled_at", "datetime", "When the label was determined (UTC)"],
            ["confidence", "float", "Label confidence 0.0–1.0 (default 1.0)"],
        ],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 4 — SCORING PIPELINE
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Chapter 4: Scoring Pipeline", level=1)

    doc.add_heading("4.1 Request Flow", level=2)
    doc.add_paragraph(
        "The scoring pipeline is orchestrated by ScoringOrchestrator in scoring/orchestrator.py. "
        "The full request flow is:"
    )
    add_code_block([
        "API Request",
        "  → Validate (Pydantic strict mode)",
        "  → Assemble Features (Redis pipelines, behavioral profiles)",
        "  → Rules Engine (Tier 1, <1ms)",
        "    → HARD BLOCK? → return BLOCK immediately",
        "  → Behavioral Engine (velocity, trust, similarity, novelty)",
        "  → ML Model Scoring (Phase-based: Cold Start / Semi / Supervised)",
        "  → Policy Floor: max(model_score, heuristic_score)",
        "  → Soft Rule Boost: +risk_boost if triggered",
        "  → Decision Mapping: <0.40→APPROVE, 0.40-0.85→REVIEW, >=0.85→BLOCK",
        "  → Async: GNN scoring, SHAP explanation (for REVIEW band)",
        "  → Emit audit event (Kafka + Redis cache)",
        "  → Return ScoringResponse",
    ])

    doc.add_heading("4.2 Three-Phase Model Lifecycle", level=2)

    doc.add_heading("Phase 1: Cold Start (Unsupervised)", level=3)
    doc.add_paragraph(
        "No labeled data required. Uses an ensemble of three unsupervised models:"
    )
    add_bullet("VAE (Variational Autoencoder) — Reconstruction error anomaly score (weight: 0.40)")
    add_bullet("Isolation Forest — Path length anomaly score (weight: 0.35)")
    add_bullet("Empirical Tail Detector — Robust z-score per feature (weight: 0.25)")
    doc.add_paragraph("Gating criteria to advance to Phase 2:")
    add_bullet("Minimum 500 fraud labels (configurable via PHASE1_MIN_FRAUD_LABELS)")
    add_bullet("Minimum 500,000 transactions")
    add_bullet("Minimum 8 weeks of data")
    add_bullet("PR-AUC >= 0.65 on held-out evaluation")

    doc.add_heading("Phase 2: Semi-Supervised", level=3)
    doc.add_paragraph(
        "Bridges unsupervised and supervised learning using pseudo-labeling:"
    )
    add_bullet("Pseudo-labels generated from Phase 1 scores (high confidence only)")
    add_bullet("XGBoost trained on confirmed labels + pseudo-labels (3:1 weight ratio)")
    add_bullet("CalibratedClassifierCV with isotonic calibration")
    add_bullet("Blended score: w_cold * cold_score + w_xgb * xgb_proba")
    add_bullet("Graph-based label propagation (account-device-merchant-IP)")
    doc.add_paragraph("Gating criteria to advance to Phase 3:")
    add_bullet("Minimum 5,000 fraud labels")
    add_bullet("PR-AUC >= 0.78")

    doc.add_heading("Phase 3: Supervised", level=3)
    doc.add_paragraph(
        "Full supervised ensemble with champion-challenger architecture:"
    )
    add_bullet("Champion: CatBoost (production inference)")
    add_bullet("Challengers: XGBoost, LightGBM, FT-Transformer, TabNet (offline evaluation)")
    add_bullet("Stacked ensemble: XGBoost + LightGBM + CatBoost → Logistic Regression meta-learner")
    add_bullet("5-fold out-of-fold stacking for meta-features")
    add_bullet("SMOTEENN for class imbalance handling")
    add_bullet("Isotonic + Temperature scaling (blended 0.7/0.3) for calibration")
    add_bullet("Conformal prediction for uncertainty quantification")
    add_bullet("SHAP explanations for REVIEW band decisions (sync)")

    doc.add_heading("4.3 Champion-Challenger Evaluation", level=2)
    doc.add_paragraph(
        "The Champion-Challenger framework in models/supervised/evaluator.py continuously "
        "evaluates challenger models against the production champion:"
    )
    add_bullet("Metrics: PR-AUC, ROC-AUC, F2-score, Precision, Recall, F1, Accuracy, FPR, Calibration Error, Brier Score")
    add_bullet("Promotion criteria: Min PR-AUC improvement of 0.01, max FPR of 0.01, max calibration error of 0.05, max latency ratio of 2.0x")
    add_bullet("Auto-approval with rollback support (max 5 rollback versions)")
    add_bullet("Version pinning: model_version, training_hash, feature_hash, dataset_hash for reproducibility")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 5 — RULES ENGINE
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Chapter 5: Rules Engine", level=1)

    doc.add_heading("5.1 Rule Types", level=2)
    add_table(
        ["Type", "Description", "Example"],
        [
            ["blocklist", "Hard block if entity is in Redis set", "BLOCKLIST_ACCOUNT — block known fraud accounts"],
            ["threshold", "Compare a feature value against a threshold", "AMOUNT_HARD_CAP — block amounts > 1,000,000"],
            ["expression", "Safe AST-based expression evaluation", "NEW_ACCT_HIGH_VALUE — new account + high amount"],
            ["velocity", "Entity velocity over time window", "VELOCITY_SPIKE_1M — >10 txns in 1 minute"],
            ["geo", "Geographic anomaly detection", "IMPOSSIBLE_TRAVEL — geo velocity > 900 km/h"],
        ],
    )

    doc.add_heading("5.2 Rule Configuration", level=2)
    doc.add_paragraph(
        "Rules are defined in config/rules.yaml with hot-reload support. "
        "Each rule has the following structure:"
    )
    add_code_block([
        "rules:",
        "  - id: VELOCITY_SPIKE_1M",
        '    description: "Velocity spike: > 10 transactions in 1 minute"',
        "    type: velocity",
        '    action: hard_block        # or soft_boost',
        '    severity: high            # critical, high, medium, low',
        "    velocity:",
        "      entity_type: account",
        "      window_seconds: 60",
        "      threshold: 10",
        '      operator: ">"',
    ])

    doc.add_heading("5.3 Rule Evaluation Flow", level=2)
    doc.add_paragraph(
        "Rules are evaluated in config/rule order. The RulesEngine in scoring/rules_engine.py "
        "processes each rule sequentially:"
    )
    add_bullet("Blocklist rules: Redis SISMEMBER check against ft:blocklist:{list_name} sets")
    add_bullet("Threshold rules: Feature comparison with configurable operators (>, >=, <, <=, ==, !=)")
    add_bullet("Expression rules: Safe AST parser (no eval/exec), supports comparisons, boolean ops, arithmetic")
    add_bullet("Velocity rules: Maps entity type to velocity feature keys (acct_v_1m_count, dev_v_5m_count)")
    add_bullet("Geo rules: Checks impossible_travel, cross_border, sanctioned_country features")

    doc.add_paragraph()
    doc.add_paragraph("Rule actions:")
    add_bullet("hard_block — Immediately returns BLOCK with score=1.0 (no ML scoring)")
    add_bullet("soft_boost — Adds a configurable boost to the ML score (capped at 0.30)")

    doc.add_heading("5.4 Blocklist Management", level=2)
    doc.add_paragraph(
        "Blocklists are managed via the Admin API (/v1/admin/blocklist) with full audit trail, "
        "TTL support, and tenant isolation. Blocklist entries are stored in Redis with metadata "
        "(added_by, reason, expiry) and logged to an audit trail."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 6 — MACHINE LEARNING MODELS
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Chapter 6: Machine Learning Models", level=1)

    doc.add_heading("6.1 Cold Start: Isolation Forest + VAE", level=2)
    doc.add_paragraph(
        "Implemented in models/cold_start/ensemble.py. The ColdStartEnsemble combines three "
        "unsupervised models that require no labeled data:"
    )
    add_table(
        ["Model", "Algorithm", "Weight", "Signal"],
        [
            ["VAE", "Variational Autoencoder (PyTorch)", "0.40", "Reconstruction error — anomalies have high error"],
            ["Isolation Forest", "scikit-learn IsolationForest", "0.35", "Path length — anomalies have short paths"],
            ["Tail Detector", "Empirical tail (GPD)", "0.25", "Robust z-score per feature — tail probability"],
        ],
    )
    doc.add_paragraph(
        "The ensemble uses weighted combination with score calibration via training percentiles. "
        "Explainability is provided through per-component contribution breakdowns."
    )

    doc.add_heading("6.2 Semi-Supervised: Ensemble Methods", level=2)
    doc.add_paragraph(
        "Implemented in models/supervised/semi_supervised.py. The SemiSupervisedBridge "
        "uses pseudo-labeling to bootstrap a supervised model from unsupervised scores:"
    )
    add_bullet("Pseudo-labels: High-confidence predictions from Phase 1 (score > 0.9 or < 0.1)")
    add_bullet("XGBoost training: Confirmed labels (weight=3) + pseudo-labels (weight=1)")
    add_bullet("Calibration: CalibratedClassifierCV with isotonic regression")
    add_bullet("Blended scoring: w_cold * cold_score + w_xgb * xgb_proba")

    doc.add_heading("6.3 Supervised: CatBoost Champion + Challengers", level=2)

    doc.add_heading("Champion Model (CatBoost)", level=3)
    doc.add_paragraph(
        "Implemented in models/supervised/champion.py. The ChampionModel is a single CatBoost "
        "classifier optimized for production inference:"
    )
    add_bullet("Native categorical feature handling (no target encoding leakage)")
    add_bullet("Built-in class imbalance handling (auto_class_weights: Balanced)")
    add_bullet("Hyperparameters: iterations=1000, depth=6, learning_rate=0.05, subsample=0.8")
    add_bullet("Early stopping after 50 rounds with AUC evaluation metric")
    add_bullet("GPU support for training acceleration")

    doc.add_heading("Challenger Models", level=3)
    doc.add_paragraph(
        "Implemented in models/supervised/challengers.py. Challengers are trained offline "
        "and evaluated continuously — they are NEVER used in production inference:"
    )
    add_table(
        ["Challenger", "Key Hyperparameters"],
        [
            ["XGBoost", "n_estimators=500, max_depth=6, learning_rate=0.05, subsample=0.8"],
            ["LightGBM", "n_estimators=500, max_depth=6, num_leaves=63, is_unbalance=True"],
            ["FT-Transformer", "n_epochs=100, batch_size=256, d_token=64, n_heads=4"],
            ["TabNet", "n_d=64, n_a=64, n_steps=3, gamma=1.3, max_epochs=100"],
        ],
    )

    doc.add_heading("6.4 Model Calibration", level=2)
    doc.add_paragraph(
        "Probability calibration ensures predicted probabilities match observed frequencies:"
    )
    add_bullet("Isotonic Regression — Non-parametric, fits an isotonic function (preferred)")
    add_bullet("Platt Scaling — Logistic regression on log-odds (simpler, parametric)")
    add_bullet("Temperature Scaling — Single parameter rescaling (used in ensemble: 0.7 isotonic + 0.3 temperature)")
    add_bullet("Conformal Prediction — Provides prediction sets for uncertainty quantification")

    doc.add_heading("6.5 Model Registry and Promotion", level=2)
    doc.add_paragraph(
        "Models are tracked via MLflow and the internal ModelRegistry (scoring/orchestrator.py):"
    )
    add_bullet("Artifact storage: artifacts/models/{tenant_id}/phase{1,2,3}/")
    add_bullet("Hot-reload: Watchdog filesystem observer detects model file changes")
    add_bullet("Atomic swap: Double-buffered staging → active with RLock for zero-downtime reloads")
    add_bullet("Version pinning: model_version, training_hash, feature_hash, dataset_hash")
    add_bullet("Warmup: Dummy inference on load to prime JIT compilation and caches")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 7 — BEHAVIORAL INTELLIGENCE
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Chapter 7: Behavioral Intelligence", level=1)

    doc.add_heading("7.1 Customer Profiles", level=2)
    doc.add_paragraph(
        "Implemented in behavior/profiles/customer.py. CustomerBehaviorProfile maintains "
        "online statistics for each customer using Welford's algorithm:"
    )
    add_bullet("Transaction statistics: total count, total amount, amount EMA")
    add_bullet("Amount statistics: Online mean/variance, rolling window percentiles")
    add_bullet("Time patterns: Preferred hours, preferred weekdays (frequency counts)")
    add_bullet("Merchant preferences: Merchant frequency, preferred merchants")
    add_bullet("Country/geo preferences: Country frequency distribution")
    add_bullet("Device tracking: Device frequency, trusted devices set, fingerprint tracking")
    add_bullet("IP tracking: IP frequency, known IPs set")

    doc.add_heading("7.2 Merchant Profiles", level=2)
    doc.add_paragraph(
        "Implemented in behavior/profiles/merchant.py. MerchantBehaviorProfile tracks "
        "merchant-level patterns:"
    )
    add_bullet("Amount statistics: Online mean/variance for transaction amounts")
    add_bullet("Transaction patterns: Volume trends, peak hours")
    add_bullet("Risk signals: Chargeback rate, fraud rate at merchant level")

    doc.add_heading("7.3 Device and Beneficiary Profiles", level=2)
    doc.add_paragraph(
        "Additional behavioral profiles in behavior/profiles/:"
    )
    add_bullet("DeviceBehaviorProfile — Device-level transaction patterns, shared account tracking")
    add_bullet("BeneficiaryBehaviorProfile — Beneficiary (destination account) patterns and trust scores")
    add_bullet("PaymentInstrumentProfile — Payment instrument usage patterns")

    doc.add_heading("7.4 Feature Generation Modules", level=2)
    doc.add_paragraph(
        "Behavioral features are generated in real-time by modules in behavior/feature_generation/:"
    )
    add_table(
        ["Module", "Key Features"],
        [
            ["velocity.py", "acct_v_1m_count, acct_v_1h_count, acct_v_24h_count, acct_v_24h_total_amt"],
            ["trust.py", "device_trust_score, merchant_trust_score, beneficiary_trust_score, customer_reputation"],
            ["similarity.py", "merchant_similarity, device_similarity, country_similarity, typing_similarity"],
            ["novelty.py", "is_new_device, is_new_merchant, is_new_country, is_new_ip"],
            ["generator.py", "Orchestrates all feature modules into unified feature vector"],
        ],
    )

    doc.add_heading("7.5 Online Statistics (Welford's Algorithm)", level=2)
    doc.add_paragraph(
        "The OnlineMeanVariance class in behavior/utils/online_statistics.py implements "
        "Welford's online algorithm for numerically stable single-pass mean and variance computation:"
    )
    add_code_block([
        "class OnlineMeanVariance:",
        "    count: int = 0",
        "    mean: float = 0.0",
        "    m2: float = 0.0  # Sum of squared differences from mean",
        "",
        "    def update(self, value: float):",
        "        self.count += 1",
        "        delta = value - self.mean",
        "        self.mean += delta / self.count",
        "        delta2 = value - self.mean",
        "        self.m2 += delta * delta2",
        "",
        "    @property",
        "    def variance(self) -> float:",
        "        return self.m2 / (self.count - 1) if self.count >= 2 else 0.0",
        "",
        "    def get_zscore(self, value: float) -> float:",
        "        return (value - self.mean) / self.std if self.std > 0 else 0.0",
    ])

    doc.add_paragraph(
        "Additional utilities: ExponentialMovingAverage, RollingWindow, CountMinSketch, "
        "and haversine_distance for geographic calculations."
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 8 — DASHBOARD
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Chapter 8: Dashboard", level=1)
    doc.add_paragraph(
        "The Streamlit dashboard in dashboard/ provides seven main pages for monitoring "
        "and managing the fraud detection platform."
    )

    doc.add_heading("8.1 Overview Page", level=2)
    doc.add_paragraph(
        "dashboard/pages/overview.py — High-level KPIs and transaction volume charts."
    )
    add_bullet("KPI metrics: Total transactions (90d), fraud rate, blocked count, in-review count, avg latency")
    add_bullet("24-hour transaction volume and fraud rate dual-axis chart")
    add_bullet("Decision distribution pie chart (APPROVE / REVIEW / BLOCK)")
    add_bullet("Tenant selector for multi-tenant viewing")

    doc.add_heading("8.2 EDA Page", level=2)
    doc.add_paragraph(
        "dashboard/pages/eda.py — Exploratory Data Analysis for feature distributions."
    )
    add_bullet("Feature distribution histograms with fraud/legit overlay")
    add_bullet("Correlation heatmap for numeric features")
    add_bullet("Data quality metrics (missing values, outliers)")
    add_bullet("Transaction type and channel breakdowns")

    doc.add_heading("8.3 Model Performance Page", level=2)
    doc.add_paragraph(
        "dashboard/pages/model_performance.py — Model evaluation metrics and comparisons."
    )
    add_bullet("PR-AUC and ROC-AUC curves")
    add_bullet("Confusion matrix visualization")
    add_bullet("Precision-Recall trade-off at different thresholds")
    add_bullet("Model version comparison (champion vs challengers)")

    doc.add_heading("8.4 Explainability Page", level=2)
    doc.add_paragraph(
        "dashboard/pages/explainability.py — SHAP-based model explanations."
    )
    add_bullet("Feature importance rankings (global)")
    add_bullet("SHAP waterfall plots for individual transactions")
    add_bullet("Rule contribution breakdown for rules-engine decisions")
    add_bullet("Component breakdown for cold-start model (VAE/IF/Tail)")

    doc.add_heading("8.5 Live Monitoring Page", level=2)
    doc.add_paragraph(
        "dashboard/pages/live_monitoring.py — Real-time scoring stream."
    )
    add_bullet("Live transaction feed with scores and decisions")
    add_bullet("Latency percentiles (P50, P95, P99) over time")
    add_bullet("Scoring throughput (transactions/second)")
    add_bullet("Error rate and SLA compliance")

    doc.add_heading("8.6 Drift Detection Page", level=2)
    doc.add_paragraph(
        "dashboard/pages/drift.py — Feature and model drift monitoring."
    )
    add_bullet("PSI (Population Stability Index) per feature over time")
    add_bullet("KL divergence trends")
    add_bullet("Feature importance drift")
    add_bullet("Concept drift detection (label rate changes)")

    doc.add_heading("8.7 Compliance Page", level=2)
    doc.add_paragraph(
        "dashboard/pages/compliance.py — Regulatory compliance views."
    )
    add_bullet("Audit trail for all scoring decisions")
    add_bullet("Model explainability reports for regulators")
    add_bullet("Data retention and GDPR compliance metrics")
    add_bullet("PCI DSS compliance indicators (no raw card data stored)")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 9 — API REFERENCE
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Chapter 9: API Reference", level=1)

    doc.add_heading("9.1 Authentication", level=2)
    doc.add_paragraph(
        "FraudTrap uses API Key authentication via the X-API-Key header. "
        "Admin endpoints additionally support Bearer Token (JWT) authentication. "
        "Rate limiting is enforced at 1000 requests per minute per tenant using slowapi."
    )

    doc.add_heading("9.2 Scoring Endpoints", level=2)
    add_table(
        ["Method", "Endpoint", "Description"],
        [
            ["POST", "/v1/score", "Score a single transaction for fraud risk. P95 target < 100ms."],
            ["POST", "/v1/score/batch", "Score up to 1000 transactions in a single request."],
            ["GET", "/v1/recent", "Retrieve recent scored transactions (dashboard cache)."],
            ["GET", "/v1/explain/{trace_id}", "Retrieve SHAP explanation by trace ID."],
        ],
    )

    doc.add_heading("9.3 Label Ingestion", level=2)
    add_table(
        ["Method", "Endpoint", "Description"],
        [
            ["POST", "/v1/labels", "Ingest ground-truth fraud labels from chargeback or manual review."],
        ],
    )
    doc.add_paragraph(
        "Labels are written to Kafka topic fraudtrap.labels.incoming for asynchronous "
        "consumption by the label worker and training pipeline."
    )

    doc.add_heading("9.4 Operations Endpoints", level=2)
    add_table(
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/health", "Health check with version and environment info."],
            ["GET", "/metrics", "Prometheus metrics (request count, latency histogram)."],
            ["GET", "/v1/phase/{tenant_id}", "Current model phase and version for a tenant."],
            ["GET", "/v1/drift/{tenant_id}", "Real-time drift metrics (PSI) by feature."],
            ["GET", "/v1/lifecycle/{tenant_id}", "Lifecycle metrics: label counts, phase progression, scoring history."],
        ],
    )

    doc.add_heading("9.5 Admin Endpoints", level=2)
    add_table(
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/v1/admin/rules", "List all active rules."],
            ["POST", "/v1/admin/rules", "Create or update a rule."],
            ["DELETE", "/v1/admin/rules/{rule_id}", "Delete a rule."],
            ["POST", "/v1/admin/rules/reload", "Force reload rules from file/Redis."],
            ["POST", "/v1/admin/blocklist", "Add entry to blocklist with TTL and audit."],
            ["DELETE", "/v1/admin/blocklist/{list_name}/{value}", "Remove entry from blocklist."],
            ["GET", "/v1/admin/blocklist/{list_name}", "List all entries in a blocklist."],
            ["GET", "/v1/admin/blocklist/audit/{tenant_id}", "Get blocklist audit log."],
            ["POST", "/v1/admin/models/reload", "Force reload all models from disk."],
            ["GET", "/v1/admin/models/status", "Get model loading status and versions."],
            ["GET", "/v1/admin/health", "Admin health check with component status."],
        ],
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 10 — DEPLOYMENT
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Chapter 10: Deployment", level=1)

    doc.add_heading("10.1 Docker Compose Setup", level=2)
    doc.add_paragraph(
        "FraudTrap is deployed via docker/docker-compose.yml. Start the stack with:"
    )
    add_code_block([
        "cd C:\\Users\\Tommie-YV\\Downloads\\fraudtrap",
        "docker compose -f docker\\docker-compose.yml up -d api dashboard live_simulator",
    ])
    doc.add_paragraph("Check service status:")
    add_code_block([
        "docker compose -f docker\\docker-compose.yml ps",
    ])

    doc.add_heading("10.2 Environment Variables", level=2)
    add_table(
        ["Variable", "Default", "Description"],
        [
            ["ENVIRONMENT", "development", "Runtime environment (development, production, testing)"],
            ["LOG_LEVEL", "INFO", "Logging level"],
            ["API_HOST", "0.0.0.0", "API bind address"],
            ["API_PORT", "8000", "API listen port"],
            ["API_WORKERS", "4", "Number of API worker processes"],
            ["SCORING_TIMEOUT_MS", "90", "Hard scoring timeout in milliseconds"],
            ["MODEL_DIR", "artifacts/models", "Path to model artifacts"],
            ["KAFKA_BROKERS", "localhost:9092", "Kafka broker list"],
            ["REDIS_HOST", "localhost", "Redis host"],
            ["REDIS_PORT", "6379", "Redis port"],
            ["CLICKHOUSE_HOST", "localhost", "ClickHouse host"],
            ["CLICKHOUSE_PORT", "9000", "ClickHouse native port"],
            ["POSTGRES_URL", "postgresql://...", "PostgreSQL connection URL"],
            ["MLFLOW_TRACKING_URI", "http://localhost:5000", "MLflow tracking server URL"],
        ],
    )

    doc.add_heading("10.3 Health Checks", level=2)
    doc.add_paragraph(
        "Every service in docker-compose.yml includes health checks. "
        "The API service exposes a /health endpoint that returns version and environment info. "
        "The admin endpoint /v1/admin/health checks Redis, ClickHouse, and PostgreSQL connectivity."
    )
    add_code_block([
        "GET /health",
        '{',
        '  "status": "ok",',
        '  "version": "1.0.0",',
        '  "environment": "production",',
        '  "timestamp": "2025-01-15T14:23:11Z"',
        '}',
    ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # CHAPTER 11 — OPERATIONS
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Chapter 11: Operations", level=1)

    doc.add_heading("11.1 Monitoring", level=2)
    doc.add_paragraph(
        "FraudTrap collects metrics via the monitoring/ module:"
    )
    add_bullet("Metrics Collector (monitoring/metrics_collector.py) — Aggregates scoring metrics into ClickHouse rollups")
    add_bullet("Request latency (P50, P95, P99) via Prometheus Histogram")
    add_bullet("Decision distribution (APPROVE/REVIEW/BLOCK counts)")
    add_bullet("Fraud capture rate")
    add_bullet("Model PR-AUC / ROC-AUC")
    add_bullet("Feature freshness (time since last Redis key update)")
    add_bullet("Behavioral fallback rates (cold-start vs profile-based scoring)")

    doc.add_heading("11.2 Alerting", level=2)
    doc.add_paragraph(
        "Structured alerting via monitoring/alerts.py with PagerDuty and Slack integration:"
    )
    add_table(
        ["Alert Rule", "Condition", "Severity", "Channel"],
        [
            ["SLA Breach", "P95 > 90ms for 5 minutes", "Critical", "PagerDuty"],
            ["Drift Spike", "PSI > 0.25 for any feature", "Warning", "Slack"],
            ["Concept Drift", "Label rate change > 20%", "Warning", "Slack"],
            ["Performance Drop", "PR-AUC drop > 5%", "Critical", "PagerDuty"],
            ["Data Quality", "> 10% features zero-valued", "Warning", "Slack"],
            ["Scoring Errors", "5xx rate > 1%", "Critical", "PagerDuty"],
            ["Behavioral Fallback", "Cold-start > 50%", "Warning", "Slack"],
        ],
    )
    doc.add_paragraph(
        "Alert deduplication: 15-minute cooldown per alert type. "
        "Runbook links are auto-attached to alert payloads."
    )

    doc.add_heading("11.3 Model Retraining", level=2)
    doc.add_paragraph(
        "Automated retraining is configured in config/supervised.yaml:"
    )
    add_bullet("Default schedule: Weekly on Sunday at 02:00 UTC (cron: 0 2 * * 0)")
    add_bullet("Training pipeline: training/pipeline.py — dataset builder + model training")
    add_bullet("Data generation: scripts/generate_sample_data.py — synthetic labeled data")
    add_bullet("Live training: scripts/run_training.py — triggered by label_worker when gating criteria met")
    add_bullet("Label lag: 70-day buffer for chargeback arrival (configurable via LABEL_LAG_DAYS)")
    add_bullet("Training window: 180 days of labeled data (configurable via TRAINING_WINDOW_DAYS)")

    doc.add_heading("11.4 Drift Detection", level=2)
    doc.add_paragraph(
        "Real-time drift monitoring via monitoring/drift.py:"
    )
    add_bullet("PSI (Population Stability Index): Compares baseline vs current feature distributions")
    add_bullet("KL Divergence: Measures distribution shift for continuous features")
    add_bullet("Mean/Std Shift: Tracks changes in feature statistics over time")
    add_bullet("Embedding Drift: Centroid distance analysis for latent representations")
    add_bullet("Concept Drift: Label rate and prediction rate change detection")
    add_bullet("Thresholds: PSI > 0.20 triggers drift alert (configurable via PSI_DRIFT_THRESHOLD)")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # APPENDIX A — CONFIGURATION REFERENCE
    # ═══════════════════════════════════════════════════════════════════════════
    doc.add_heading("Appendix A: Configuration Reference", level=1)

    doc.add_heading("A.1 settings.py Environment Variables", level=2)
    add_table(
        ["Setting", "Type", "Default", "Description"],
        [
            ["app_name", "str", "FraudTrap", "Application name"],
            ["app_version", "str", "1.0.0", "Application version"],
            ["environment", "str", "development", "Runtime environment"],
            ["api_host", "str", "0.0.0.0", "API bind address"],
            ["api_port", "int", "8000", "API port"],
            ["api_workers", "int", "4", "Number of workers"],
            ["scoring_timeout_ms", "int", "90", "Hard scoring timeout (ms)"],
            ["model_dir", "str", "artifacts/models", "Model artifacts directory"],
            ["kafka_brokers", "str", "localhost:9092", "Kafka brokers"],
            ["redis_host", "str", "localhost", "Redis host"],
            ["redis_port", "int", "6379", "Redis port"],
            ["clickhouse_host", "str", "localhost", "ClickHouse host"],
            ["postgres_url", "str", "postgresql://...", "PostgreSQL URL"],
            ["phase1_min_fraud_labels", "int", "500", "Cold Start → Semi gate"],
            ["phase2_min_fraud_labels", "int", "5000", "Semi → Supervised gate"],
            ["phase1_min_pr_auc", "float", "0.65", "Phase 1 PR-AUC threshold"],
            ["phase2_min_pr_auc", "float", "0.78", "Phase 2 PR-AUC threshold"],
            ["score_block_threshold", "float", "0.85", "Auto-block above this score"],
            ["score_review_low", "float", "0.40", "Review band lower bound"],
            ["score_review_high", "float", "0.85", "Review band upper bound"],
            ["psi_drift_threshold", "float", "0.20", "PSI drift alert threshold"],
            ["dp_epsilon", "float", "10.0", "Differential privacy budget"],
        ],
    )

    doc.add_heading("A.2 rules.yaml Structure", level=2)
    doc.add_paragraph("Top-level structure of config/rules.yaml:")
    add_code_block([
        "rules:",
        "  - id: <rule_id>",
        '    description: "<human-readable description>"',
        "    type: <blocklist|threshold|expression|velocity|geo>",
        "    action: <hard_block|soft_boost>",
        "    severity: <critical|high|medium|low>",
        "    enabled: true|false",
        "    boost: <float>        # only for soft_boost actions",
        "    max_boost: <float>    # cap for accumulated boost",
        "    # Type-specific config:",
        "    blocklist:             # for type=blocklist",
        "      entity: <account|device|ip|merchant|country>",
        "      list_name: <redis_set_name>",
        "    threshold:             # for type=threshold",
        "      feature: <feature_name>",
        "      operator: <|<=|>|>=|==|!=>",
        "      threshold: <float>",
        "    expression:            # for type=expression",
        "      expression: <safe_python_expression>",
        "    velocity:              # for type=velocity",
        "      entity_type: <account|device|ip>",
        "      window_seconds: <int>",
        "      threshold: <float>",
        "      operator: <|<=|>|=>",
        "    geo:                   # for type=geo",
        "      rule_type: <impossible_travel|cross_border|sanctioned_country>",
        "      threshold_kmh: <float>",
        "",
        "global:",
        "  default_boost: 0.05",
        "  max_boost: 0.30",
        "  enabled: true",
    ])

    doc.add_heading("A.3 supervised.yaml Structure", level=2)
    doc.add_paragraph("Top-level structure of config/supervised.yaml:")
    add_code_block([
        "champion:",
        "  algorithm: catboost",
        "  params:",
        "    iterations: 1000",
        "    depth: 6",
        "    learning_rate: 0.05",
        "    loss_function: Logloss",
        "    eval_metric: AUC",
        "    auto_class_weights: Balanced",
        "    early_stopping_rounds: 50",
        "  calibration:",
        "    method: isotonic",
        "    enabled: true",
        "",
        "challengers:",
        "  algorithms: [xgboost, lightgbm, ft_transformer, tabnet]",
        "  xgboost:",
        "    n_estimators: 500",
        "    max_depth: 6",
        "    learning_rate: 0.05",
        "  lightgbm:",
        "    n_estimators: 500",
        "    num_leaves: 63",
        "    is_unbalance: true",
        "",
        "evaluation:",
        "  metrics: [pr_auc, roc_auc, f2_score, precision, recall, f1]",
        "  threshold_analysis:",
        "    enabled: true",
        "  latency:",
        "    enabled: true",
        "",
        "promotion:",
        "  criteria:",
        "    min_pr_auc_improvement: 0.01",
        "    max_fpr: 0.01",
        "    max_calibration_error: 0.05",
        "    max_latency_ratio: 2.0",
        "",
        "training_pipeline:",
        "  output_dir: models/supervised/training_runs",
        "  model_storage:",
        "    base_dir: models/supervised/saved_models",
        "    compression: true",
        "  scheduling:",
        "    cron: '0 2 * * 0'  # Weekly Sunday 02:00 UTC",
    ])

    # ═══════════════════════════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════════════════════════
    output_path = Path(r"C:\Users\Tommie-YV\Downloads\fraudtrap\FraudTrap_Complete_Study_Guide.docx")
    doc.save(str(output_path))
    print(f"[+] Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    ensure_python_docx()
    build_document()
